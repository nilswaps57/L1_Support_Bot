"""DOCX fallback parser using python-docx."""

from __future__ import annotations

import io
from typing import Any

from l1_support_bot.domain.errors import ProcessingError
from l1_support_bot.domain.models.document import FileType
from l1_support_bot.domain.models.parsed_document import DocumentElement, ParsedDocument

try:
    from docx import Document as WordDocument
except ImportError:  # pragma: no cover
    WordDocument = None  # type: ignore[assignment]


class PythonDocxParser:
    async def parse(self, content: bytes, file_type: FileType) -> ParsedDocument:
        return self._parse(content, file_type)

    def _parse(self, content: bytes, file_type: FileType) -> ParsedDocument:
        if file_type is not FileType.DOCX:
            raise ProcessingError("The DOCX fallback cannot process this document type.")
        if WordDocument is None:
            raise ProcessingError(
                "The DOCX fallback parser is unavailable.", details={"stage": "parsing"}
            )
        try:
            document = WordDocument(io.BytesIO(content))
            elements: list[DocumentElement] = []
            section_path: tuple[str, ...] = ()
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                style = paragraph.style.name.lower() if paragraph.style else ""
                if "heading" in style:
                    level = (
                        int(style.rsplit(" ", 1)[-1]) if style.rsplit(" ", 1)[-1].isdigit() else 1
                    )
                    section_path = section_path[: level - 1] + (text,)
                    elements.append(
                        DocumentElement(
                            "heading", text, heading_level=level, section_path=section_path
                        )
                    )
                else:
                    elements.append(DocumentElement("paragraph", text, section_path=section_path))
            for table in document.tables:
                rows = tuple(tuple(cell.text.strip() for cell in row.cells) for row in table.rows)
                table_text = "\n".join(" | ".join(row) for row in rows)
                if table_text:
                    elements.append(
                        DocumentElement(
                            "table", table_text, section_path=section_path, table_rows=rows
                        )
                    )
            if not elements:
                raise ProcessingError(
                    "The DOCX did not contain readable content.", details={"stage": "parsing"}
                )
            return ParsedDocument(tuple(elements), "docx")
        except ProcessingError:
            raise
        except Exception as exc:
            raise ProcessingError(
                "The DOCX could not be read.", details={"stage": "parsing"}
            ) from exc

    @staticmethod
    def make_test_docx(heading: str, cell: str) -> bytes:
        if WordDocument is None:
            raise RuntimeError("python-docx is required for DOCX fixtures")
        document: Any = WordDocument()
        document.add_heading(heading, level=1)
        document.add_paragraph("Procedure content")
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = cell
        output = io.BytesIO()
        document.save(output)
        return output.getvalue()
